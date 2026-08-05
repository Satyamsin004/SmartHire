import re
import io
import uuid
import logging
from typing import Dict, Any, List, Optional
from fastapi import HTTPException, status
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select

from app.services.ai_engine import ai_engine
from app.models.domain import (
    Candidate, Resume, ResumeSkill, ResumeEducation, ResumeExperience,
    ResumeInternship, ResumeProject, ResumeCertification, ResumeAchievement,
    ResumeLanguage, ResumeATS, Notification
)

logger = logging.getLogger(__name__)

class ResumeService:
    def extract_text_from_file_bytes(self, file_bytes: bytes, filename: str, max_size_mb: int = 10) -> str:
        """Validates file size/format and extracts clean plain text from PDF, DOCX, or TXT."""
        # 1. Size Validation
        if len(file_bytes) > max_size_mb * 1024 * 1024:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File size exceeds maximum allowed limit of {max_size_mb}MB."
            )

        if not file_bytes or len(file_bytes) == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Uploaded file is empty or corrupted."
            )

        ext = filename.split(".")[-1].lower() if "." in filename else ""
        raw_text = ""

        # 2. PDF Extraction
        if ext == "pdf":
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            raw_text += t + "\n"
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}. Trying fallback pypdf...")
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            raw_text += t + "\n"
                except Exception as ex:
                    logger.error(f"PDF extraction error: {ex}")

        # 3. DOCX Extraction
        elif ext in ["docx", "doc"]:
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                full_text = []
                for para in doc.paragraphs:
                    if para.text:
                        full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            full_text.append(row_text)
                raw_text = "\n".join(full_text)
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")

        # 4. TXT / Plain Text Fallback
        if not raw_text or not raw_text.strip():
            try:
                raw_text = file_bytes.decode("utf-8", errors="ignore")
            except Exception:
                raw_text = ""

        # 5. Clean Extracted Text
        clean_text = self._clean_text(raw_text)

        if not clean_text or len(clean_text.strip()) < 20:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not extract readable text from file. Please ensure the document is not an image-only scan or encrypted."
            )

        return clean_text

    def _clean_text(self, text: str) -> str:
        """Cleans null bytes, normalizes whitespace and removes unprintable characters."""
        if not text:
            return ""
        text = text.replace('\x00', '')
        text = re.sub(r'[\r\n\t]+', '\n', text)
        text = re.sub(r' +', ' ', text)
        return text.strip()

    async def parse_and_store_resume(
        self,
        db: AsyncSession,
        candidate: Candidate,
        file_name: str,
        file_path: str,
        raw_text: str
    ) -> Dict[str, Any]:
        """Runs Gemini 12-section parsing, calculates ATS telemetry, normalizes database schema, handles versioning and updates Candidate profile."""
        
        # 1. Gemini Structured Parsing
        parsed_data = await ai_engine.parse_resume_to_json(raw_text)

        # Extract sub-blocks safely and normalize list types
        p_info = parsed_data.get("personal_information") or {}
        p_summary = parsed_data.get("professional_summary") or {}

        def _to_list(val):
            if isinstance(val, list): return val
            if isinstance(val, dict): return [val]
            return []

        work_exp = _to_list(parsed_data.get("work_experience"))
        internships = _to_list(parsed_data.get("internships"))
        projects = _to_list(parsed_data.get("projects"))
        education = _to_list(parsed_data.get("education"))
        tech_skills = _to_list(parsed_data.get("technical_skills"))
        soft_skills = _to_list(parsed_data.get("soft_skills"))
        certs = _to_list(parsed_data.get("certifications"))
        achievements = _to_list(parsed_data.get("achievements"))
        languages = _to_list(parsed_data.get("languages"))
        ats_data = parsed_data.get("ats_analysis") or {}

        # 2. Manage Versioning: Fetch past versions
        res_existing = await db.execute(
            select(Resume).where(Resume.candidate_id == candidate.id).order_by(Resume.version.desc())
        )
        existing_resumes = res_existing.scalars().all()
        
        highest_version = 0
        for r in existing_resumes:
            if r.version and r.version > highest_version:
                highest_version = r.version
            r.is_active = False # Deactivate previous versions
        
        new_version = highest_version + 1

        # 3. Auto-update Candidate Profile
        if p_info.get("phone") and p_info.get("phone") != "Not Available":
            candidate.phone = p_info["phone"]
        if p_info.get("location") and p_info.get("location") != "Not Available":
            candidate.location = p_info["location"]
        if p_info.get("github") and p_info.get("github") != "Not Available":
            candidate.github_url = p_info["github"]
        if p_info.get("linkedin") and p_info.get("linkedin") != "Not Available":
            candidate.linkedin_url = p_info["linkedin"]
        if p_info.get("portfolio") and p_info.get("portfolio") != "Not Available":
            candidate.portfolio_url = p_info["portfolio"]
        if p_summary.get("summary") and p_summary.get("summary") != "Not Available":
            candidate.bio = p_summary["summary"]
        if p_summary.get("experience_years") and p_summary.get("experience_years") != "Not Available":
            candidate.experience_level = p_summary["experience_years"]
        candidate.resume_url = file_path

        # 4. Create Active Resume Record in PostgreSQL
        ats_score_raw = ats_data.get("ats_score", 85.0)
        try:
            ats_score_val = float(ats_score_raw)
        except (ValueError, TypeError):
            ats_score_val = 85.0

        deg_raw = (education[0].get("degree") or "").lower() if (education and isinstance(education[0], dict)) else ""
        if "bachelor" in deg_raw or "b.s" in deg_raw or "b.e" in deg_raw or "b.tech" in deg_raw or "bs" in deg_raw:
            edu_level = "Bachelor's Degree"
        elif "master" in deg_raw or "m.s" in deg_raw or "m.e" in deg_raw or "m.tech" in deg_raw or "ms" in deg_raw:
            edu_level = "Master's Degree"
        elif "phd" in deg_raw or "doctor" in deg_raw:
            edu_level = "Doctorate / Ph.D."
        else:
            edu_level = education[0].get("degree") if (education and isinstance(education[0], dict)) else "Bachelor's Degree"

        resume = Resume(
            id=f"res-{uuid.uuid4().hex[:8]}",
            candidate_id=candidate.id,
            file_name=file_name,
            file_path=file_path,
            raw_text=raw_text,
            summary=p_summary.get("summary"),
            objective=p_summary.get("objective"),
            ats_score=ats_score_val,
            keyword_density=ats_data.get("repeated_skills", {}) if isinstance(ats_data.get("repeated_skills"), dict) else {},
            missing_skills=ats_data.get("missing_keywords", []) if isinstance(ats_data.get("missing_keywords"), list) else [],
            projects=[p.get("project_name") if isinstance(p, dict) else str(p) for p in projects],
            certifications=[c.get("certificate_name") if isinstance(c, dict) else str(c) for c in certs],
            languages=[l.get("language_name") if isinstance(l, dict) else str(l) for l in languages],
            experience_years=p_summary.get("experience_years"),
            education_level=edu_level,
            version=new_version,
            is_active=True
        )
        db.add(resume)
        await db.flush()

        # 5. Insert Relational Education
        for edu in education:
            if isinstance(edu, dict):
                db.add(ResumeEducation(
                    resume_id=resume.id,
                    degree=edu.get("degree"),
                    college=edu.get("college"),
                    university=edu.get("university"),
                    board=edu.get("board"),
                    cgpa=str(edu.get("cgpa")) if edu.get("cgpa") else None,
                    percentage=str(edu.get("percentage")) if edu.get("percentage") else None,
                    year=str(edu.get("year")) if edu.get("year") else None,
                    branch=edu.get("branch"),
                    specialization=edu.get("specialization")
                ))

        # 6. Insert Relational Experience
        for exp in work_exp:
            if isinstance(exp, dict):
                db.add(ResumeExperience(
                    resume_id=resume.id,
                    company_name=exp.get("company_name"),
                    job_title=exp.get("job_title"),
                    employment_type=exp.get("employment_type", "Full-Time"),
                    location=exp.get("location"),
                    joining_date=exp.get("joining_date"),
                    ending_date=exp.get("ending_date"),
                    is_current=exp.get("is_current", False),
                    duration=exp.get("duration"),
                    responsibilities=exp.get("responsibilities", []),
                    achievements=exp.get("achievements", []),
                    technologies=exp.get("technologies", []),
                    projects_worked=exp.get("projects_worked", [])
                ))

        # 7. Insert Relational Internships
        for intern in internships:
            if isinstance(intern, dict):
                db.add(ResumeInternship(
                    resume_id=resume.id,
                    company=intern.get("company"),
                    role=intern.get("role"),
                    duration=intern.get("duration"),
                    description=intern.get("description"),
                    skills_used=intern.get("skills_used", []),
                    projects=intern.get("projects", []),
                    technologies=intern.get("technologies", [])
                ))

        # 8. Insert Relational Projects
        for proj in projects:
            if isinstance(proj, dict):
                db.add(ResumeProject(
                    resume_id=resume.id,
                    project_name=proj.get("project_name", "Untitled Project"),
                    description=proj.get("description"),
                    role=proj.get("role"),
                    responsibilities=proj.get("responsibilities", []),
                    technologies=proj.get("technologies", []),
                    programming_languages=proj.get("programming_languages", []),
                    frameworks=proj.get("frameworks", []),
                    database=proj.get("database"),
                    cloud=proj.get("cloud"),
                    github_link=proj.get("github_link"),
                    live_link=proj.get("live_link"),
                    achievements=proj.get("achievements", [])
                ))

        # 9. Insert Relational Technical & Soft Skills
        for sk in tech_skills:
            if isinstance(sk, dict):
                db.add(ResumeSkill(
                    resume_id=resume.id,
                    skill_name=sk.get("skill_name", "Skill"),
                    category=sk.get("category", "Technical"),
                    proficiency=sk.get("proficiency", "Expert")
                ))

        for ssk in soft_skills:
            if isinstance(ssk, dict):
                db.add(ResumeSkill(
                    resume_id=resume.id,
                    skill_name=ssk.get("skill_name", "Soft Skill"),
                    category="Soft",
                    proficiency=ssk.get("proficiency", "Proficient")
                ))

        # 10. Insert Certifications
        for c in certs:
            if isinstance(c, dict):
                db.add(ResumeCertification(
                    resume_id=resume.id,
                    certificate_name=c.get("certificate_name", "Certificate"),
                    organization=c.get("organization"),
                    issue_date=c.get("issue_date"),
                    credential_id=c.get("credential_id"),
                    verification_url=c.get("verification_url")
                ))

        # 11. Insert Achievements
        for ach in achievements:
            if isinstance(ach, dict):
                db.add(ResumeAchievement(
                    resume_id=resume.id,
                    title=ach.get("title", "Achievement"),
                    category=ach.get("category", "Award"),
                    description=ach.get("description")
                ))

        # 12. Insert Languages
        for lang in languages:
            if isinstance(lang, dict):
                db.add(ResumeLanguage(
                    resume_id=resume.id,
                    language_name=lang.get("language_name", "Language"),
                    proficiency=lang.get("proficiency", "Fluent")
                ))

        # 13. Insert ATS & Keywords Record
        km_raw = ats_data.get("keyword_match_percentage", 80.0)
        try:
            km_val = float(km_raw)
        except (ValueError, TypeError):
            km_val = 80.0

        db.add(ResumeATS(
            resume_id=resume.id,
            ats_score=ats_score_val,
            keyword_match_percentage=km_val,
            technical_keywords=ats_data.get("technical_keywords", []),
            domain_keywords=ats_data.get("domain_keywords", []),
            missing_keywords=ats_data.get("missing_keywords", []),
            repeated_skills=ats_data.get("repeated_skills", {}),
            strengths=ats_data.get("strengths", []),
            weaknesses=ats_data.get("weaknesses", []),
            formatting_issues=ats_data.get("formatting_issues", []),
            suggestions=ats_data.get("suggestions", [])
        ))

        # 14. Create Notification for Candidate
        notif = Notification(
            user_id=candidate.user_id,
            title="Resume Parsed Successfully!",
            message=f"Your resume '{file_name}' has been parsed and stored in PostgreSQL with an ATS score of {ats_score_val}%.",
            notification_type="resume_parsed"
        )
        db.add(notif)

        await db.commit()
        logger.info(f"Resume version v{new_version} successfully stored in PostgreSQL for Candidate {candidate.id}")

        return await self.get_full_parsed_resume(db, resume.id)

    async def get_full_parsed_resume(self, db: AsyncSession, resume_id: str) -> Dict[str, Any]:
        """Retrieves active normalized resume + all child relational entities from PostgreSQL."""
        res_r = await db.execute(select(Resume).where(Resume.id == resume_id))
        resume = res_r.scalar_one_or_none()
        if not resume:
            return {}

        res_c = await db.execute(select(Candidate).where(Candidate.id == resume.candidate_id))
        candidate = res_c.scalar_one_or_none()

        # Load relational entities
        res_edu = await db.execute(select(ResumeEducation).where(ResumeEducation.resume_id == resume.id))
        educations = res_edu.scalars().all()

        res_exp = await db.execute(select(ResumeExperience).where(ResumeExperience.resume_id == resume.id))
        experiences = res_exp.scalars().all()

        res_intern = await db.execute(select(ResumeInternship).where(ResumeInternship.resume_id == resume.id))
        internships = res_intern.scalars().all()

        res_proj = await db.execute(select(ResumeProject).where(ResumeProject.resume_id == resume.id))
        projects = res_proj.scalars().all()

        res_sk = await db.execute(select(ResumeSkill).where(ResumeSkill.resume_id == resume.id))
        skills = res_sk.scalars().all()

        res_cert = await db.execute(select(ResumeCertification).where(ResumeCertification.resume_id == resume.id))
        certs = res_cert.scalars().all()

        res_ach = await db.execute(select(ResumeAchievement).where(ResumeAchievement.resume_id == resume.id))
        achievements = res_ach.scalars().all()

        res_lang = await db.execute(select(ResumeLanguage).where(ResumeLanguage.resume_id == resume.id))
        languages = res_lang.scalars().all()

        res_ats = await db.execute(select(ResumeATS).where(ResumeATS.resume_id == resume.id))
        ats = res_ats.scalar_one_or_none()

        full_name = candidate.user.full_name if (candidate and candidate.user) else "Candidate"
        email = candidate.user.email if (candidate and candidate.user) else "Not Available"
        phone = candidate.phone if candidate else "Not Available"
        location = candidate.location if candidate else "Not Available"
        github = candidate.github_url if candidate else "Not Available"
        linkedin = candidate.linkedin_url if candidate else "Not Available"
        portfolio = candidate.portfolio_url if candidate else "Not Available"
        summary_val = resume.summary or "Not Available"
        exp_years_val = resume.experience_years or "Not Available"
        edu_level_val = resume.education_level or (educations[0].degree if educations else "Not Available")

        skills_clean = [s.skill_name for s in skills]
        certs_clean = [c.certificate_name for c in certs]
        projects_clean = [p.project_name for p in projects]
        langs_clean = [l.language_name for l in languages]

        return {
            "status": "success",
            "message": f"Resume v{resume.version} retrieved successfully.",
            "resume_id": resume.id,
            "candidate_id": resume.candidate_id,
            "version": resume.version,
            "is_active": resume.is_active,
            "file_name": resume.file_name,
            "file_path": resume.file_path,
            "ats_score": resume.ats_score,
            "candidate_name": full_name,
            "name": full_name,
            "email": email,
            "phone": phone,
            "location": location,
            "github": github,
            "linkedin": linkedin,
            "portfolio": portfolio,
            "summary": summary_val,
            "experience_years": exp_years_val,
            "education_level": edu_level_val,
            "personal_information": {
                "full_name": full_name,
                "email": email,
                "phone": phone,
                "location": location,
                "github": github,
                "linkedin": linkedin,
                "portfolio": portfolio
            },
            "professional_summary": {
                "summary": summary_val,
                "objective": resume.objective or "Not Available",
                "experience_years": exp_years_val
            },
            "work_experience": [
                {
                    "company_name": e.company_name,
                    "job_title": e.job_title,
                    "employment_type": e.employment_type,
                    "location": e.location,
                    "joining_date": e.joining_date,
                    "ending_date": e.ending_date,
                    "is_current": e.is_current,
                    "duration": e.duration,
                    "responsibilities": e.responsibilities or [],
                    "achievements": e.achievements or [],
                    "technologies": e.technologies or []
                } for e in experiences
            ],
            "internships": [
                {
                    "company": i.company,
                    "role": i.role,
                    "duration": i.duration,
                    "description": i.description,
                    "skills_used": i.skills_used or [],
                    "technologies": i.technologies or []
                } for i in internships
            ],
            "projects": [
                {
                    "project_name": p.project_name,
                    "description": p.description,
                    "role": p.role,
                    "responsibilities": p.responsibilities or [],
                    "technologies": p.technologies or [],
                    "github_link": p.github_link,
                    "live_link": p.live_link
                } for p in projects
            ],
            "education": [
                {
                    "degree": edu.degree,
                    "college": edu.college,
                    "university": edu.university,
                    "cgpa": edu.cgpa,
                    "percentage": edu.percentage,
                    "year": edu.year,
                    "branch": edu.branch
                } for edu in educations
            ],
            "skills": [
                {
                    "skill_name": s.skill_name,
                    "category": s.category,
                    "proficiency": s.proficiency
                } for s in skills
            ],
            "certifications": [
                {
                    "certificate_name": c.certificate_name,
                    "organization": c.organization,
                    "issue_date": c.issue_date,
                    "credential_id": c.credential_id
                } for c in certs
            ],
            "achievements": [
                {
                    "title": a.title,
                    "category": a.category,
                    "description": a.description
                } for a in achievements
            ],
            "languages": [
                {
                    "language_name": l.language_name,
                    "proficiency": l.proficiency
                } for l in languages
            ],
            "ats_analysis": {
                "ats_score": ats.ats_score if ats else resume.ats_score,
                "keyword_match_percentage": ats.keyword_match_percentage if ats else 80.0,
                "technical_keywords": ats.technical_keywords if ats else [],
                "domain_keywords": ats.domain_keywords if ats else [],
                "missing_keywords": ats.missing_keywords if ats else resume.missing_skills,
                "strengths": ats.strengths if ats else [],
                "weaknesses": ats.weaknesses if ats else [],
                "formatting_issues": ats.formatting_issues if ats else [],
                "suggestions": ats.suggestions if ats else []
            },
            "resume": {
                "id": resume.id,
                "version": resume.version,
                "file_name": resume.file_name,
                "file_path": resume.file_path,
                "summary": summary_val,
                "skills": skills_clean,
                "experience_years": exp_years_val,
                "education_level": edu_level_val,
                "projects": projects_clean,
                "certifications": certs_clean,
                "languages": langs_clean
            }
        }

    async def get_resume_versions(self, db: AsyncSession, candidate_id: str) -> List[Dict[str, Any]]:
        """Retrieves list of all resume upload versions for candidate history."""
        res_r = await db.execute(
            select(Resume).where(Resume.candidate_id == candidate_id).order_by(Resume.version.desc())
        )
        resumes = res_r.scalars().all()

        return [
            {
                "id": r.id,
                "version": r.version,
                "is_active": r.is_active,
                "file_name": r.file_name,
                "file_path": r.file_path,
                "ats_score": r.ats_score,
                "created_at": r.created_at.isoformat() if r.created_at else None
            } for r in resumes
        ]

    async def match_job_description(
        self,
        candidate_skills: List[str],
        job_description: str,
        required_skills: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Calculates real ATS Match score, matching skills, missing skills, and qualitative feedback based on skills and role alignment."""
        if not candidate_skills:
            candidate_skills = []

        jd_text = (job_description or "").lower()
        req_skills = [s.strip() for s in (required_skills or []) if s and s.strip()]

        matching_skills = []
        missing_skills = []

        # If required skills provided, evaluate candidate against required skills
        if req_skills:
            for skill in req_skills:
                sk_low = skill.lower()
                is_matched = any(sk_low in cs.lower() or cs.lower() in sk_low for cs in candidate_skills)
                if is_matched:
                    if skill not in matching_skills:
                        matching_skills.append(skill)
                else:
                    if skill not in missing_skills:
                        missing_skills.append(skill)
        else:
            # Fallback: extract technical keywords from JD
            common_tech = [
                "python", "fastapi", "react", "typescript", "postgresql", "docker",
                "kubernetes", "aws", "sql", "java", "c++", "pytorch", "tensorflow",
                "machine learning", "ai", "node.js", "graphql", "redis"
            ]
            extracted_reqs = [t.title() for t in common_tech if t in jd_text]
            for skill in extracted_reqs:
                sk_low = skill.lower()
                is_matched = any(sk_low in cs.lower() or cs.lower() in sk_low for cs in candidate_skills)
                if is_matched:
                    if skill not in matching_skills:
                        matching_skills.append(skill)
                else:
                    if skill not in missing_skills:
                        missing_skills.append(skill)

        total_targets = len(matching_skills) + len(missing_skills)
        if total_targets > 0:
            match_ratio = len(matching_skills) / total_targets
            raw_score = match_ratio * 100.0
        else:
            raw_score = 0.0

        match_score = round(min(100.0, max(0.0, raw_score)), 1)
        recommendation = "Shortlist" if match_score >= 80.0 else "Reject"

        logger.info("ATS Engine Evaluated: match_score=%.1f%% (%d/%d skills matched) -> %s", match_score, len(matching_skills), total_targets, recommendation)

        return {
            "match_score": match_score,
            "matching_skills": matching_skills,
            "missing_skills": missing_skills,
            "analysis": {
                "ats_score": match_score,
                "matched_skills": matching_skills,
                "missing_skills": missing_skills,
                "recommendation": recommendation,
                "reasoning": f"Matched {len(matching_skills)} of {total_targets} required requisition skills."
            }
        }

resume_service = ResumeService()

